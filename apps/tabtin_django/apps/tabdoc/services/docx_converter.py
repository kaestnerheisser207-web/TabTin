"""
ProseMirror JSON → Word (DOCX) 转换器

将 TipTap/ProseMirror 编辑器的 JSON 文档结构转为 python-docx Word 文档。
支持的节点类型：heading, paragraph, blockquote, codeBlock, bulletList,
orderedList, taskList, table, horizontalRule, image, mathematics。

超链接使用 python-docx 的 OPC 低层 API 实现真实可点击链接。
图片节点通过 ThreadPoolExecutor 并行下载后嵌入 DOCX，下载失败时 fallback 到占位文字。
"""
from __future__ import annotations

import io
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed, wait, FIRST_COMPLETED
from typing import Any

try:
    import requests
except ImportError:
    requests = None  # type: ignore[assignment]

logger = logging.getLogger(__name__)

_IMAGE_SAFE_SCHEMES = ("https://", "http://")
_LINK_SAFE_SCHEMES = ("http://", "https://", "mailto:", "tel:")
# python-docx 可直接嵌入的光栅 MIME（不含 SVG；SVG 需先栅格化）
_IMAGE_ALLOWED_CONTENT_TYPES = (
    "image/png", "image/jpeg", "image/gif", "image/webp", "image/jpg",
    "image/bmp", "image/x-ms-bmp", "image/tiff", "image/avif",
    "image/heic", "image/heif", "image/apng",
)
# 下载阶段额外允许 SVG：导出前会栅格化为 PNG（与 TabSlide PPTX 对齐）
_IMAGE_DOWNLOAD_CONTENT_TYPES = _IMAGE_ALLOWED_CONTENT_TYPES + ("image/svg+xml",)
_PER_IMAGE_TIMEOUT = 10
_TOTAL_IMAGE_BUDGET = 30  # BE-18: 降低总时间预算（并行后无需 60s）
_MAX_IMAGE_BYTES = 10 * 1024 * 1024  # BE-18: 单张图片 10 MB
_MAX_TOTAL_IMAGE_BYTES = 50 * 1024 * 1024  # BE-18: 总内存限制 50 MB
_MAX_IMAGE_WORKERS = 5  # BE-18: 并行下载线程数
_PNG_MAGIC = b"\x89PNG\r\n\x1a\n"
# 与 TabSlide PPTX 默认 svg_scale=2.0 对齐；1x 栅格化后拉大显示会发糊
_DOCX_SVG_SUPERSAMPLE = 2.0


def _looks_like_svg_bytes(img_bytes: bytes) -> bool:
    head = img_bytes[:512].lstrip().lower()
    return head.startswith(b"<svg") or b"<svg" in head


def _normalize_image_bytes_for_docx(
    image_bytes: bytes,
    src_hint: str | None = None,
    target_width_px: int | None = None,
    target_height_px: int | None = None,
) -> bytes:
    """将 SVG / 非常见格式尽量转为 python-docx 可嵌入的 PNG。

    复用 TabSlide PPTX 导出的归一化实现；转换失败时返回原始 bytes，
    由调用方检测后降级为占位符（避免把 SVG 塞进 add_picture）。
    """
    if not image_bytes:
        return image_bytes
    # 已是 PNG 则跳过
    if image_bytes.startswith(_PNG_MAGIC):
        return image_bytes
    try:
        from apps.tabslide.services.pptx_io import _normalize_image_bytes_for_pptx
    except Exception as exc:
        logger.debug("TabSlide image normalize unavailable: %s", exc)
        return image_bytes
    try:
        return _normalize_image_bytes_for_pptx(
            image_bytes,
            src_hint=src_hint,
            target_width_px=float(target_width_px) if target_width_px else None,
            target_height_px=float(target_height_px) if target_height_px else None,
            supersample=_DOCX_SVG_SUPERSAMPLE,
        )
    except Exception as exc:
        logger.warning("DOCX image normalize failed: %s", exc)
        return image_bytes


def _decode_data_uri_image(src: str) -> tuple[str, bytes]:
    """解码 data URI 图片，支持 base64 与 charset 百分号编码（对话拖入 SVG 常用后者）。

    返回 (mime, raw_bytes)。拒绝非法 / 超大 payload。
    """
    import base64 as _b64
    from urllib.parse import unquote_to_bytes

    parts = src.split(",", 1)
    if len(parts) != 2:
        raise ValueError("Invalid data URI: missing comma")
    header_part = parts[0]
    header_lower = header_part.lower()
    if not header_lower.startswith("data:"):
        raise ValueError("Invalid data URI: missing data: scheme")

    mime = header_lower.replace("data:", "", 1).split(";")[0].strip()
    if not mime:
        raise ValueError("Invalid data URI: missing MIME type")

    payload = parts[1]
    if ";base64" in header_lower:
        estimated_bytes = len(payload) * 3 // 4
        if estimated_bytes > _MAX_IMAGE_BYTES:
            raise ValueError(
                f"Data URI 估算大小 {estimated_bytes} 字节超过上限 "
                f"{_MAX_IMAGE_BYTES} 字节，拒绝解码"
            )
        img_bytes = _b64.b64decode(payload)
    else:
        # percent-encoded（如 data:image/svg+xml;charset=utf-8,%3Csvg...）
        estimated_bytes = len(payload)
        if estimated_bytes > _MAX_IMAGE_BYTES * 2:
            raise ValueError(
                f"Data URI payload 过长 ({estimated_bytes})，拒绝解码"
            )
        img_bytes = unquote_to_bytes(payload)

    if len(img_bytes) > _MAX_IMAGE_BYTES:
        raise ValueError(f"Data URI image too large: {len(img_bytes)} bytes")
    if not img_bytes:
        raise ValueError("Data URI image is empty")
    return mime, img_bytes


def _prepare_image_bytes_for_docx(
    img_bytes: bytes,
    *,
    src_hint: str | None = None,
    mime: str | None = None,
    width_px: int | None = None,
    height_px: int | None = None,
) -> bytes | None:
    """归一化并校验：成功返回可嵌入 bytes，失败返回 None（调用方写占位符）。"""
    hint = src_hint or (f"data:{mime};base64," if mime else None)
    normalized = _normalize_image_bytes_for_docx(
        img_bytes,
        src_hint=hint,
        target_width_px=width_px,
        target_height_px=height_px,
    )
    if not normalized:
        return None
    # SVG 未转成功时不要交给 python-docx（会抛错或写坏包）
    if _looks_like_svg_bytes(normalized):
        logger.debug("SVG rasterize failed for DOCX embed; falling back to placeholder")
        return None
    return normalized


def _safe_int(value: Any) -> int | None:
    """安全转换为正整数，无效值返回 None。"""
    if value is None:
        return None
    try:
        n = int(str(value).replace("px", "").replace("pt", "").strip())
        return n if n > 0 else None
    except (ValueError, TypeError):
        return None


def _resolve_and_validate_ip(hostname: str) -> str | None:
    """解析 DNS 并验证所有返回 IP 均为公网地址，防止 SSRF。

    返回第一个安全的 IP 字符串，若任一地址为内网/保留地址或解析失败则返回 None。
    调用方必须使用返回的 IP 直接连接，禁止让 HTTP 库重新解析 DNS（防 DNS 重绑定）。
    """
    import ipaddress
    import socket
    try:
        infos = socket.getaddrinfo(hostname, None, socket.AF_UNSPEC, socket.SOCK_STREAM)
        if not infos:
            return None
        first_ip: str | None = None
        for _, _, _, _, sockaddr in infos:
            ip = ipaddress.ip_address(sockaddr[0])
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
                return None
            if first_ip is None:
                first_ip = sockaddr[0]
        return first_ip
    except (socket.gaierror, ValueError, OSError):
        return None


def _platform_asset_object_key(src: str) -> str | None:
    """Resolve Muse-owned asset refs without treating arbitrary URLs as OSS keys."""
    try:
        from apps.services.oss.services.public_assets import public_asset_object_key_from_ref
        return public_asset_object_key_from_ref(src)
    except Exception as exc:
        logger.debug("Platform asset ref resolution failed for %s: %s", src[:80], exc)
        return None


def _download_platform_image(src: str) -> bytes | None:
    object_key = _platform_asset_object_key(src)
    if not object_key:
        return None
    try:
        from apps.services.oss.services.factory import get_oss_service
        result = get_oss_service().download_file(object_key)
        if not isinstance(result, dict):
            return None
        if not result.get("success"):
            return None
        data = result.get("data") or {}
        content = data.get("content")
        if not isinstance(content, (bytes, bytearray)) or not content:
            return None
        if len(content) > _MAX_IMAGE_BYTES:
            return None
        content_type = (data.get("content_type") or "").split(";")[0].strip().lower()
        if content_type and content_type not in _IMAGE_DOWNLOAD_CONTENT_TYPES:
            return None
        return bytes(content)
    except Exception as exc:
        logger.debug("Platform image download failed for %s: %s", src[:80], exc)
        return None


def _pinned_get(src: str, pinned_ip: str, hostname: str, **kwargs) -> requests.Response:
    """使用已验证 IP 发起 HTTP GET，防止 DNS 重绑定 SSRF。

    旧实现把 URL host 换成 IP 后未设置 TLS SNI / assert_hostname，
    导致外链 PNG/JPEG（Seedream TOS 等）证书校验失败、DOCX 丢图。
    现统一走 ``ssrf_safe_request``（含正确 SNI + 证书主机名校验）。

    ``pinned_ip`` / ``hostname`` 保留以兼容既有单测 mock 签名；
    实际解析与校验由 ``ssrf_safe_request`` 完成。
    """
    del pinned_ip, hostname  # 由 ssrf_safe_request 内部 resolve-validate-pin
    from apps.services.common.url_security import ssrf_safe_request

    allow_redirects = bool(kwargs.pop("allow_redirects", False))
    timeout = kwargs.pop("timeout", _PER_IMAGE_TIMEOUT)
    return ssrf_safe_request(
        "GET",
        src,
        allow_redirects=allow_redirects,
        timeout=timeout,
        **kwargs,
    )


def image_bytes_to_data_uri(img_bytes: bytes, content_type: str | None = None) -> str:
    """将图片字节编码为 data URI（PDF 导出内联用）。"""
    import base64 as _b64

    mime = (content_type or "").split(";")[0].strip().lower()
    if mime not in _IMAGE_DOWNLOAD_CONTENT_TYPES:
        if img_bytes.startswith(_PNG_MAGIC):
            mime = "image/png"
        elif img_bytes[:3] == b"\xff\xd8\xff":
            mime = "image/jpeg"
        elif img_bytes[:6] in (b"GIF87a", b"GIF89a"):
            mime = "image/gif"
        elif img_bytes[:4] == b"RIFF" and img_bytes[8:12] == b"WEBP":
            mime = "image/webp"
        elif _looks_like_svg_bytes(img_bytes):
            mime = "image/svg+xml"
        else:
            mime = "image/png"
    encoded = _b64.b64encode(img_bytes).decode("ascii")
    return f"data:{mime};base64,{encoded}"


def _collect_image_urls(nodes: list[dict[str, Any]]) -> list[str]:
    """递归收集所有可预下载图片引用（去重，保持顺序）。"""
    urls: list[str] = []
    seen: set[str] = set()
    def _walk(node_list: list) -> None:
        for node in node_list:
            if not isinstance(node, dict):
                continue
            if node.get("type") == "image":
                src = str((node.get("attrs") or {}).get("src") or "")
                if (
                    src
                    and not src.lower().startswith("data:")
                    and (
                        any(src.lower().startswith(s) for s in _IMAGE_SAFE_SCHEMES)
                        or _platform_asset_object_key(src) is not None
                    )
                    and src not in seen
                ):
                    seen.add(src)
                    urls.append(src)
            children = node.get("content", [])
            if isinstance(children, list):
                _walk(children)
    _walk(nodes)
    return urls


def _download_single_image(src: str) -> tuple[str, bytes | None]:
    """下载单张图片，返回 (url, bytes|None)。线程安全。"""
    from urllib.parse import urlparse

    if requests is None:
        return (src, _download_platform_image(src))
    try:
        platform_bytes = _download_platform_image(src)
        if platform_bytes:
            return (src, platform_bytes)

        parsed = urlparse(src)
        hostname = parsed.hostname or ""
        if not hostname:
            return (src, None)

        # PAR-010: 先解析 DNS 并验证 IP，然后使用该 IP 直接连接，
        # 防止 DNS 重绑定攻击（TOCTOU 漏洞）
        pinned_ip = _resolve_and_validate_ip(hostname)
        if pinned_ip is None:
            return (src, None)

        resp = _pinned_get(
            src, pinned_ip, hostname,
            timeout=_PER_IMAGE_TIMEOUT, stream=True,
            allow_redirects=False,
        )
        try:
            if resp.status_code in (301, 302, 303, 307, 308):
                return (src, None)

            content_length = int(resp.headers.get("Content-Length", 0) or 0)
            if content_length > _MAX_IMAGE_BYTES:
                return (src, None)

            content_type = (resp.headers.get("Content-Type") or "").split(";")[0].strip().lower()
            if resp.status_code != 200 or content_type not in _IMAGE_DOWNLOAD_CONTENT_TYPES:
                return (src, None)

            chunks: list[bytes] = []
            total = 0
            for chunk in resp.iter_content(chunk_size=8192):
                total += len(chunk)
                if total > _MAX_IMAGE_BYTES:
                    return (src, None)
                chunks.append(chunk)
            return (src, b"".join(chunks))
        finally:
            resp.close()
    except Exception as exc:
        logger.debug("Image download failed for %s: %s", src[:80], exc)
        return (src, None)


def _batch_download_images(urls: list[str]) -> dict[str, bytes]:
    """BE-18: 并行下载所有图片，带总内存限制。"""
    if not urls or requests is None:
        return {}

    result: dict[str, bytes] = {}
    total_bytes = 0
    budget_exceeded = False

    import time as _time

    with ThreadPoolExecutor(max_workers=_MAX_IMAGE_WORKERS) as pool:
        futures = {pool.submit(_download_single_image, url): url for url in urls}
        deadline = _time.monotonic() + _TOTAL_IMAGE_BUDGET
        remaining = set(futures.keys())

        while remaining:
            time_left = deadline - _time.monotonic()
            if time_left <= 0:
                logger.warning(
                    "Total image download budget (%ds) exceeded, skipping %d remaining images",
                    _TOTAL_IMAGE_BUDGET,
                    len(remaining),
                )
                break

            done, remaining = wait(remaining, timeout=time_left, return_when=FIRST_COMPLETED)
            for future in done:
                if budget_exceeded:
                    continue
                try:
                    src, img_bytes = future.result(timeout=0)
                    if img_bytes and len(img_bytes) > 0:
                        if total_bytes + len(img_bytes) > _MAX_TOTAL_IMAGE_BYTES:
                            logger.warning(
                                "Total image memory budget exceeded (%d bytes), skipping remaining images",
                                total_bytes,
                            )
                            budget_exceeded = True
                            continue
                        total_bytes += len(img_bytes)
                        result[src] = img_bytes
                except Exception as exc:
                    logger.debug("Image future failed: %s", exc)

    return result


def pm_json_to_docx_bytes(
    pm_json: dict[str, Any],
    markdown_fallback: str = "",
) -> bytes:
    """将 ProseMirror JSON 转为 DOCX 二进制。"""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError(
            "导出 Word 需要 python-docx 库，请执行: pip install python-docx"
        )

    doc = DocxDocument()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    content = pm_json.get("content", [])
    if not isinstance(content, list) or not content:
        doc.add_paragraph(markdown_fallback or "")
        return _to_bytes(doc)

    # BE-18: 先并行下载所有远程图片，再渲染
    image_urls = _collect_image_urls(content)
    image_cache = _batch_download_images(image_urls)

    renderer = _DocxRenderer(doc, image_cache=image_cache)
    for node in content:
        if isinstance(node, dict):
            renderer.render_node(node)

    return _to_bytes(doc)


def _to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _DocxRenderer:
    """递归遍历 ProseMirror 节点树，渲染到 python-docx Document。"""

    _HEADING_STYLE_MAP = {
        1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3',
        4: 'Heading 4', 5: 'Heading 5', 6: 'Heading 6',
    }

    def __init__(self, doc, image_cache: dict[str, bytes] | None = None):
        self._doc = doc
        self._image_cache = image_cache or {}

    def render_node(self, node: dict[str, Any]) -> None:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        node_type = node.get("type", "")
        children = node.get("content", [])
        child_nodes = (
            [c for c in children if isinstance(c, dict)]
            if isinstance(children, list) else []
        )

        if node_type == "heading":
            level = max(1, min(6, int((node.get("attrs") or {}).get("level", 1))))
            p = self._doc.add_paragraph(
                style=self._HEADING_STYLE_MAP.get(level, 'Heading 4')
            )
            self._add_inline_runs(p, child_nodes)

        elif node_type == "paragraph":
            p = self._doc.add_paragraph()
            self._add_inline_runs(p, child_nodes)

        elif node_type == "blockquote":
            for child in child_nodes:
                if child.get("type") == "paragraph":
                    has_quote = 'Quote' in [s.name for s in self._doc.styles]
                    p = self._doc.add_paragraph(
                        style='Quote' if has_quote else 'Normal'
                    )
                    self._add_inline_runs(p, child.get("content", []))
                else:
                    self.render_node(child)

        elif node_type == "codeBlock":
            text = _extract_text(node)
            attrs = node.get("attrs") or {}
            lang = str(attrs.get("language") or "").strip()
            p = self._doc.add_paragraph()
            if lang:
                lang_run = p.add_run(f"[{lang}]\n")
                lang_run.font.name = 'Courier New'
                lang_run.font.size = Pt(8)
                lang_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            pf = p.paragraph_format
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)

        elif node_type == "bulletList":
            self._render_list(child_nodes, 'List Bullet', depth=0)

        elif node_type == "orderedList":
            self._render_list(child_nodes, 'List Number', depth=0)

        elif node_type == "taskList":
            for item in child_nodes:
                checked = bool((item.get("attrs") or {}).get("checked"))
                prefix = "☑ " if checked else "☐ "
                p = self._doc.add_paragraph()
                run = p.add_run(prefix)
                run.font.size = Pt(11)
                self._render_list_item_blocks(p, item)

        elif node_type == "table":
            self._render_table(child_nodes)

        elif node_type == "horizontalRule":
            p = self._doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run("─" * 40)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)

        elif node_type == "image":
            attrs = node.get("attrs") or {}
            alt = attrs.get("alt", "")
            src = str(attrs.get("src", ""))
            width_px = _safe_int(attrs.get("width"))
            height_px = _safe_int(attrs.get("height"))
            self._embed_or_placeholder(src, alt, width_px=width_px, height_px=height_px)

        elif node_type == "mathematics":
            attrs = node.get("attrs") or {}
            latex = attrs.get("latex", "") or attrs.get("value", "")
            is_display = attrs.get("display", False)
            if is_display:
                # P0-4/P0-5: 块级公式拆为独立段落 ($$, 内容行…, $$)，
                # 匹配 markdown_to_pm_json 的 _BLOCK_MATH_OPEN_RE 解析格式
                math_lines = ["$$"] + latex.split("\n") + ["$$"]
                for line_text in math_lines:
                    p = self._doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run(line_text)
                    run.font.name = 'Cambria Math'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x66, 0x33, 0x99)
            else:
                p = self._doc.add_paragraph()
                run = p.add_run(f"${latex}$")
                run.font.name = 'Cambria Math'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x66, 0x33, 0x99)

        elif node_type == "youtube":
            attrs = node.get("attrs") or {}
            src = attrs.get("src", "")
            p = self._doc.add_paragraph()
            run = p.add_run(f"[YouTube: {src}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        elif node_type == "tabdataBlock":
            attrs = node.get("attrs") or {}
            title = str(attrs.get("title") or "未命名表格")
            p = self._doc.add_paragraph()
            run = p.add_run(f"[表格: {title}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        elif node_type == "htmlBlock":
            attrs = node.get("attrs") or {}
            title = str(attrs.get("title") or "未命名 HTML")
            p = self._doc.add_paragraph()
            run = p.add_run(f"[HTML: {title}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        else:
            for child in child_nodes:
                self.render_node(child)

    def _embed_or_placeholder(
        self, src: str, alt: str,
        width_px: int | None = None, height_px: int | None = None,
        target_paragraph=None,
    ) -> None:
        """尝试嵌入图片到 DOCX，失败则 fallback 到占位文字。

        BE-18: 远程图片使用预下载缓存（_image_cache），不再同步阻塞。
        PAR-043: 读取节点 width/height attrs 按比例缩放，不再强制 5.5 英寸。
        DC-4: target_paragraph 不为 None 时嵌入到指定段落（用于表格单元格）。
        SVG：先栅格化为 PNG（对话拖入的 charset data URI / OSS .svg 均可），再嵌入。
        """
        from docx.shared import RGBColor

        if src and src.lower().startswith("data:"):
            try:
                mime, raw_bytes = _decode_data_uri_image(src)
                if not mime.startswith("image/"):
                    raise ValueError(f"Unsupported image MIME type: {mime}")
                prepared = _prepare_image_bytes_for_docx(
                    raw_bytes,
                    src_hint=src,
                    mime=mime,
                    width_px=width_px,
                    height_px=height_px,
                )
                if prepared is None:
                    raise ValueError(f"Unable to prepare data URI image for DOCX: {mime}")
                self._insert_image(
                    prepared, alt,
                    width_px=width_px, height_px=height_px,
                    target_paragraph=target_paragraph,
                )
                return
            except (ValueError, TypeError, OSError) as exc:
                logger.debug("Data URI image embed failed: %s", exc)

        img_bytes = self._image_cache.get(src)
        if img_bytes and len(img_bytes) > 0:
            try:
                prepared = _prepare_image_bytes_for_docx(
                    img_bytes,
                    src_hint=src,
                    width_px=width_px,
                    height_px=height_px,
                )
                if prepared is None:
                    raise ValueError("Cached image could not be normalized for DOCX")
                self._insert_image(
                    prepared, alt,
                    width_px=width_px, height_px=height_px,
                    target_paragraph=target_paragraph,
                )
                return
            except Exception as exc:
                logger.debug("Image embed from cache failed for %s: %s", src[:80], exc)

        p = target_paragraph or self._doc.add_paragraph()
        placeholder = self._image_placeholder_text(src, alt)
        run = p.add_run(placeholder)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    @staticmethod
    def _image_placeholder_text(src: str, alt: str) -> str:
        label = f"[图片: {alt}]" if alt else "[图片]"
        source = str(src or "").strip()
        if source and not source.lower().startswith("data:"):
            return f"{label} {source}"
        return label

    _MAX_IMAGE_WIDTH_INCHES = 6.0
    _PX_PER_INCH = 96

    def _insert_image(
        self, img_bytes: bytes, alt: str,
        width_px: int | None = None, height_px: int | None = None,
        target_paragraph=None,
    ) -> None:
        """将图片字节嵌入 DOCX 文档，根据节点 attrs 按比例缩放。

        DC-4: target_paragraph 不为 None 时嵌入到指定段落（跳过独立 caption 段落）。
        """
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        p = target_paragraph or self._doc.add_paragraph()
        if target_paragraph is None:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()

        target_width = None
        if width_px and width_px > 0:
            width_in = width_px / self._PX_PER_INCH
            target_width = Inches(min(width_in, self._MAX_IMAGE_WIDTH_INCHES))
        else:
            target_width = Inches(self._MAX_IMAGE_WIDTH_INCHES)

        run.add_picture(io.BytesIO(img_bytes), width=target_width)
        if alt and target_paragraph is None:
            cap = self._doc.add_paragraph()
            cap_run = cap.add_run(alt)
            cap_run.font.italic = True
            cap_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            cap_run.font.size = Pt(9)

    def _render_list(self, items: list, style_name: str, depth: int = 0) -> None:
        from docx.shared import Inches

        for item in items:
            if not isinstance(item, dict):
                continue
            p = self._doc.add_paragraph(style=style_name)
            if depth > 0:
                p.paragraph_format.left_indent = Inches(0.5 * depth)
            self._render_list_item_blocks(p, item, depth=depth)

    def _render_list_item_blocks(self, first_p, item: dict, depth: int = 0) -> None:
        """渲染 listItem 内部的多个 block 节点，支持嵌套列表缩进。"""
        item_content = item.get("content", [])
        if not isinstance(item_content, list):
            return
        for i, block in enumerate(item_content):
            if not isinstance(block, dict):
                continue
            block_type = block.get("type", "")
            if block_type == "paragraph":
                target = first_p if i == 0 else self._doc.add_paragraph()
                self._add_inline_runs(target, block.get("content", []))
            elif block_type in ("bulletList", "orderedList"):
                nested_items = [c for c in block.get("content", []) if isinstance(c, dict)]
                style = 'List Bullet' if block_type == "bulletList" else 'List Number'
                self._render_list(nested_items, style, depth=depth + 1)
            else:
                self.render_node(block)

    def _add_inline_runs(self, paragraph, nodes: list) -> None:
        from docx.shared import Pt, RGBColor

        if not isinstance(nodes, list):
            return
        for node in nodes:
            if not isinstance(node, dict):
                continue

            if node.get("type") == "text":
                text = node.get("text", "")
                if not text:
                    continue
                marks = {
                    m["type"]: m
                    for m in (node.get("marks") or [])
                    if isinstance(m, dict) and "type" in m
                }

                link_href = ""
                link_mark = marks.get("link")
                if link_mark:
                    attrs = link_mark.get("attrs", {})
                    link_href = str(attrs.get("href", "")) if isinstance(attrs, dict) else ""

                if link_href:
                    _add_hyperlink(paragraph, link_href, text, marks)
                else:
                    run = paragraph.add_run(text)
                    _apply_marks(run, marks)

            elif node.get("type") == "hardBreak":
                paragraph.add_run("\n")
            elif node.get("type") == "image":
                attrs = node.get("attrs") or {}
                src = str(attrs.get("src", ""))
                alt = str(attrs.get("alt", ""))
                width_px = _safe_int(attrs.get("width"))
                height_px = _safe_int(attrs.get("height"))
                self._embed_or_placeholder(
                    src,
                    alt,
                    width_px=width_px,
                    height_px=height_px,
                    target_paragraph=paragraph,
                )
            else:
                self._add_inline_runs(paragraph, node.get("content", []))

    def _render_table(self, rows: list) -> None:
        if not rows:
            return

        row_nodes = [r for r in rows if isinstance(r, dict)]
        if not row_nodes:
            return

        num_rows = len(row_nodes)
        occupied: set[tuple[int, int]] = set()
        grid_info: list[tuple[int, int, int, int, dict]] = []
        max_col = 0

        for row_idx, row_node in enumerate(row_nodes):
            cells_data = row_node.get("content", [])
            if not isinstance(cells_data, list):
                continue
            col_idx = 0
            for cell_node in cells_data:
                if not isinstance(cell_node, dict):
                    continue
                while (row_idx, col_idx) in occupied:
                    col_idx += 1
                attrs = cell_node.get("attrs") or {}
                colspan = _safe_int(attrs.get("colspan")) or 1
                rowspan = _safe_int(attrs.get("rowspan")) or 1
                grid_info.append((row_idx, col_idx, colspan, rowspan, cell_node))
                for dr in range(rowspan):
                    for dc in range(colspan):
                        occupied.add((row_idx + dr, col_idx + dc))
                col_idx += colspan
            max_col = max(max_col, col_idx)

        if max_col == 0:
            max_col = 1

        table = self._doc.add_table(rows=num_rows, cols=max_col)
        table.style = 'Table Grid'

        for row_idx, col_idx, colspan, rowspan, _cell_node in grid_info:
            if colspan <= 1 and rowspan <= 1:
                continue
            end_row = min(row_idx + rowspan - 1, num_rows - 1)
            end_col = min(col_idx + colspan - 1, max_col - 1)
            top_left = table.cell(row_idx, col_idx)
            bottom_right = table.cell(end_row, end_col)
            if top_left._tc is not bottom_right._tc:
                top_left.merge(bottom_right)
                for extra_p in top_left.paragraphs[1:]:
                    extra_p._element.getparent().remove(extra_p._element)

        for row_idx, col_idx, _colspan, _rowspan, cell_node in grid_info:
            cell = table.cell(row_idx, col_idx)
            cell_content = cell_node.get("content", [])
            if isinstance(cell_content, list):
                self._render_cell_blocks(cell, cell_content)
            if cell_node.get("type") == "tableHeader":
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    def _render_cell_blocks(self, cell, blocks: list) -> None:
        """渲染表格单元格内的块级节点，支持 paragraph/codeBlock/image/list。"""
        from docx.shared import Pt, RGBColor

        block_idx = 0
        for block in blocks:
            if not isinstance(block, dict):
                continue
            block_type = block.get("type", "")
            p = cell.paragraphs[0] if block_idx == 0 else cell.add_paragraph()
            block_idx += 1

            if block_type == "paragraph":
                self._add_inline_runs(p, block.get("content", []))
            elif block_type == "codeBlock":
                code_text = _extract_text(block)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            elif block_type == "image":
                attrs = block.get("attrs") or {}
                alt_text = attrs.get("alt", "")
                src = str(attrs.get("src", ""))
                width_px = _safe_int(attrs.get("width"))
                height_px = _safe_int(attrs.get("height"))
                self._embed_or_placeholder(
                    src, alt_text, width_px=width_px, height_px=height_px,
                    target_paragraph=p,
                )
            elif block_type in ("bulletList", "orderedList"):
                items = [c for c in block.get("content", []) if isinstance(c, dict)]
                for j, li in enumerate(items):
                    if not isinstance(li, dict):
                        continue
                    li_p = p if j == 0 else cell.add_paragraph()
                    prefix = "• " if block_type == "bulletList" else f"{j + 1}. "
                    li_p.add_run(prefix)
                    li_content = li.get("content", [])
                    if isinstance(li_content, list):
                        for sub in li_content:
                            if isinstance(sub, dict) and sub.get("type") == "paragraph":
                                self._add_inline_runs(li_p, sub.get("content", []))
            else:
                self._add_inline_runs(p, block.get("content", []))


def _apply_marks(run, marks: dict[str, Any]) -> None:
    """将 ProseMirror marks 应用到 Word Run 上。"""
    from docx.shared import Pt, RGBColor

    if "bold" in marks or "strong" in marks:
        run.bold = True
    if "italic" in marks or "em" in marks:
        run.italic = True
    if "strike" in marks:
        run.font.strike = True
    if "underline" in marks:
        run.underline = True
    if "code" in marks:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    if "subscript" in marks:
        run.font.subscript = True
    if "superscript" in marks:
        run.font.superscript = True

    text_style = marks.get("textStyle")
    if text_style and isinstance(text_style, dict):
        ts_attrs = text_style.get("attrs", {})
        if isinstance(ts_attrs, dict):
            color = ts_attrs.get("color", "")
            if isinstance(color, str) and color.startswith("#") and len(color) == 7:
                try:
                    run.font.color.rgb = RGBColor(
                        int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16),
                    )
                except (ValueError, TypeError):
                    pass
            font_size = ts_attrs.get("fontSize")
            if font_size:
                try:
                    size_str = str(font_size).replace("px", "").replace("pt", "").strip()
                    run.font.size = Pt(float(size_str))
                except (ValueError, TypeError):
                    pass

    highlight_mark = marks.get("highlight")
    if highlight_mark and isinstance(highlight_mark, dict):
        try:
            from docx.enum.text import WD_COLOR_INDEX
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        except (ImportError, AttributeError):
            pass


def _is_safe_link_url(url: str) -> bool:
    """校验超链接 URL scheme 是否在白名单内，防止 javascript:/file:// 等危险 scheme。"""
    normalized = (url or "").strip().lower()
    return any(normalized.startswith(s) for s in _LINK_SAFE_SCHEMES)


def _add_hyperlink(paragraph, url: str, text: str, marks: dict[str, Any]) -> None:
    """在 Word 段落中添加真实可点击的超链接。

    python-docx 没有直接的 add_hyperlink API，需要操作 OPC XML。
    不安全的 URL scheme（如 javascript:, file://）将被替换为纯文本输出。
    DC-6: 通过 _apply_marks 应用所有 mark 类型（不再仅限 bold/italic）。
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not _is_safe_link_url(url):
        run = paragraph.add_run(text)
        _apply_marks(run, marks)
        return

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), '1A73E8')
    rPr.append(color_elem)

    u_elem = OxmlElement('w:u')
    u_elem.set(qn('w:val'), 'single')
    rPr.append(u_elem)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.set(qn('xml:space'), 'preserve')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    from docx.text.run import Run as _DocxRun
    _apply_marks(_DocxRun(new_run, paragraph), marks)


def _extract_text(node: dict[str, Any]) -> str:
    if node.get("type") == "text":
        return node.get("text", "")
    children = node.get("content", [])
    if not isinstance(children, list):
        return ""
    return "".join(_extract_text(c) for c in children if isinstance(c, dict))
