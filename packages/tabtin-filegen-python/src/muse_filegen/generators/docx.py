"""DOCX generator (python-docx).

Shares the block document model with the PDF generator: a flat list of typed
blocks (heading / paragraph / list / table).
"""

from __future__ import annotations

from typing import Any, Dict

from docx import Document

from muse_filegen.errors import SpecError
from muse_filegen.validate import as_text, optional_text, require_list, require_mapping

_SPEC_HELP = """docx spec (shared block model with pdf):
{
  "title": "Document title",            // optional, rendered as Title style
  "blocks": [
    {"type": "heading", "level": 1, "text": "Heading"},
    {"type": "paragraph", "text": "Body text..."},
    {"type": "list", "ordered": false, "items": ["a", "b"]},
    {"type": "table", "header": ["c1", "c2"], "rows": [["x", "y"]]}
  ]
}"""

_MAX_HEADING_LEVEL = 6


class DocxGenerator:
    file_type = "docx"
    extensions = (".docx",)

    def spec_help(self) -> str:
        return _SPEC_HELP

    def generate(self, spec: Dict[str, Any], out_path: str) -> None:
        document = Document()

        title = optional_text(spec, "title")
        if title:
            document.add_heading(title, level=0)

        blocks = require_list(spec.get("blocks", []), "blocks")
        for index, raw_block in enumerate(blocks):
            block = require_mapping(raw_block, f"blocks[{index}]")
            self._render_block(document, block, index)

        document.save(out_path)

    def _render_block(self, document: Document, block: Dict[str, Any], index: int) -> None:
        block_type = block.get("type")
        if block_type == "heading":
            level = _clamp_level(block.get("level", 1))
            document.add_heading(optional_text(block, "text"), level=level)
        elif block_type == "paragraph":
            document.add_paragraph(optional_text(block, "text"))
        elif block_type == "list":
            ordered = bool(block.get("ordered", False))
            style = "List Number" if ordered else "List Bullet"
            for item in require_list(block.get("items", []), f"blocks[{index}].items"):
                document.add_paragraph(as_text(item), style=style)
        elif block_type == "table":
            self._render_table(document, block, index)
        else:
            raise SpecError(f"blocks[{index}].type unsupported: {block_type!r}")

    def _render_table(self, document: Document, block: Dict[str, Any], index: int) -> None:
        header = block.get("header")
        rows = require_list(block.get("rows", []), f"blocks[{index}].rows")
        header_cells = (
            require_list(header, f"blocks[{index}].header") if header is not None else None
        )

        column_count = len(header_cells) if header_cells else _max_row_len(rows)
        if column_count == 0:
            raise SpecError(f"blocks[{index}] table has no columns")

        table = document.add_table(rows=0, cols=column_count)
        table.style = "Table Grid"

        if header_cells:
            cells = table.add_row().cells
            for col in range(column_count):
                cells[col].text = as_text(header_cells[col]) if col < len(header_cells) else ""

        for row_index, raw_row in enumerate(rows):
            row = require_list(raw_row, f"blocks[{index}].rows[{row_index}]")
            cells = table.add_row().cells
            for col in range(column_count):
                cells[col].text = as_text(row[col]) if col < len(row) else ""


_READ_HELP = """docx read 输出：
{
  "file_type": "docx",
  "paragraphs": ["段落文本", ...],
  "tables": [{"rows": [["c1", "c2"], ...]}],
  "text": "全部段落用换行拼接"
}"""


class DocxReader:
    file_type = "docx"
    extensions = (".docx",)

    def read_help(self) -> str:
        return _READ_HELP

    def read(self, path: str) -> Dict[str, Any]:
        document = Document(path)
        paragraphs = [p.text for p in document.paragraphs]
        tables = [
            {"rows": [[cell.text for cell in row.cells] for row in table.rows]}
            for table in document.tables
        ]
        return {
            "file_type": "docx",
            "paragraphs": paragraphs,
            "tables": tables,
            "text": "\n".join(paragraphs),
        }


def _clamp_level(value: Any) -> int:
    try:
        level = int(value)
    except (TypeError, ValueError):
        return 1
    return max(1, min(_MAX_HEADING_LEVEL, level))


def _max_row_len(rows: list) -> int:
    return max((len(r) for r in rows if isinstance(r, list)), default=0)
